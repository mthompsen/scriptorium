import io

from ingestion_service.parsers import parse


def test_markdown_tracks_heading_hierarchy() -> None:
    content = (
        b"# Handbook\n\nWelcome text.\n\n"
        b"## Leave\n\nPTO is 25 days.\n\nCarryover is 5 days.\n\n"
        b"## Remote work\n\nThree days per week.\n"
    )

    blocks = parse("text/markdown", content)

    assert [b.text for b in blocks] == [
        "Welcome text.",
        "PTO is 25 days.",
        "Carryover is 5 days.",
        "Three days per week.",
    ]
    assert blocks[1].headings == ["Handbook", "Leave"]
    assert blocks[3].headings == ["Handbook", "Remote work"]


def test_html_extracts_headings_and_paragraphs() -> None:
    content = (
        b"<html><body><h1>Policy</h1><p>Rule one.</p>"
        b"<h2>Scope</h2><li>All staff</li></body></html>"
    )

    blocks = parse("text/html", content)

    assert [b.text for b in blocks] == ["Rule one.", "All staff"]
    assert blocks[1].headings == ["Policy", "Scope"]


def test_plain_text_splits_paragraphs() -> None:
    blocks = parse("text/plain", b"First paragraph.\n\nSecond\nparagraph continues.")

    assert [b.text for b in blocks] == ["First paragraph.", "Second paragraph continues."]
    assert blocks[0].headings == []


def test_docx_uses_heading_styles() -> None:
    import docx

    document = docx.Document()
    document.add_heading("Security", level=1)
    document.add_paragraph("Use strong passwords.")
    buffer = io.BytesIO()
    document.save(buffer)

    blocks = parse(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        buffer.getvalue(),
    )

    assert blocks[0].text == "Use strong passwords."
    assert blocks[0].headings == ["Security"]


def test_unknown_mime_type_is_rejected() -> None:
    import pytest

    with pytest.raises(ValueError, match="no parser"):
        parse("application/zip", b"...")
